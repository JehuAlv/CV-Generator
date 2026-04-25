import io
from reportlab.lib import colors
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, Spacer, Table, TableStyle, HRFlowable, SimpleDocTemplate
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from cv_templates import TEMPLATES, LABELS_ES


def _build_pdf_story(data, scale, labels=None):
    if labels is None:
        labels = LABELS_ES

    tpl = TEMPLATES[data["theme"]]
    ACCENT = colors.HexColor(tpl["primary"])
    LIGHT_ACCENT = colors.HexColor(tpl["accent"])
    MID_GRAY = colors.HexColor("#555555")
    LIGHT_GRAY = colors.HexColor("#999999")
    WHITE = colors.white
    BLACK = colors.HexColor("#222222")

    def s(base):
        return base * scale

    page_w = 7.4 * inch
    header_align = TA_CENTER if tpl["header_align"] == "center" else TA_LEFT

    name_style = ParagraphStyle("Name", fontName="Helvetica-Bold", fontSize=s(20), leading=s(24), textColor=ACCENT, spaceAfter=0, alignment=header_align)
    contact_style = ParagraphStyle("Contact", fontName="Helvetica", fontSize=s(8.5), textColor=MID_GRAY, spaceAfter=0, alignment=header_align)
    subtitle_style = ParagraphStyle("Subtitle", fontName="Helvetica-Oblique", fontSize=s(10), textColor=LIGHT_ACCENT, spaceAfter=0, alignment=header_align)
    body_style = ParagraphStyle("Body", fontName="Helvetica", fontSize=s(8), textColor=MID_GRAY, spaceAfter=s(1), leading=s(10.5))
    bullet_style = ParagraphStyle("Bullet", fontName="Helvetica", fontSize=s(8), textColor=MID_GRAY, leftIndent=10, spaceAfter=s(0.5), leading=s(10.5))
    job_title_style = ParagraphStyle("JobTitle", fontName="Helvetica-Bold", fontSize=s(9.5), textColor=ACCENT, spaceAfter=s(0.5))
    company_style = ParagraphStyle("Company", fontName="Helvetica", fontSize=s(9), textColor=LIGHT_ACCENT, spaceAfter=s(0.5))
    date_style_r = ParagraphStyle("DateR", fontName="Helvetica-Oblique", fontSize=s(8.5), textColor=LIGHT_GRAY, alignment=TA_RIGHT)
    skill_label_s = ParagraphStyle("SkillLabel", fontName="Helvetica-Bold", fontSize=s(8.5), textColor=ACCENT, spaceAfter=s(1))
    skill_value_s = ParagraphStyle("SkillValue", fontName="Helvetica", fontSize=s(8.5), textColor=MID_GRAY, spaceAfter=s(1.5), leading=s(11))
    section_text_style = ParagraphStyle("SectionText", fontName="Helvetica-Bold", fontSize=s(10), textColor=ACCENT, spaceAfter=0, spaceBefore=s(6))
    section_bar_style = ParagraphStyle("SectionBar", fontName="Helvetica-Bold", fontSize=s(9), textColor=WHITE, leftIndent=4, spaceAfter=s(3), spaceBefore=s(6))

    def section_header(title):
        ss = tpl["section_style"]
        elements = []

        if ss == "bar":
            tbl = Table([[Paragraph(title, section_bar_style)]], colWidths=[page_w])
            tbl.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), ACCENT),
                ("TOPPADDING", (0, 0), (-1, -1), s(2)),
                ("BOTTOMPADDING", (0, 0), (-1, -1), s(2)),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ]))
            elements.append(tbl)

        elif ss == "text_underline":
            elements.append(Paragraph(title, section_text_style))
            elements.append(Spacer(1, s(2)))
            elements.append(HRFlowable(width="100%", thickness=1.5, color=ACCENT))

        elif ss == "left_border":
            inner_style = ParagraphStyle("LB", fontName="Helvetica-Bold", fontSize=s(9.5), textColor=ACCENT, spaceAfter=0)
            tbl = Table([[Paragraph(title, inner_style)]], colWidths=[page_w])
            tbl.setStyle(TableStyle([
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), s(3)),
                ("BOTTOMPADDING", (0, 0), (-1, -1), s(3)),
                ("LINEBEFORESTARTX", (0, 0), (0, -1), 3, ACCENT),
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F5F8FC")),
            ]))
            elements.append(Spacer(1, s(4)))
            elements.append(tbl)

        elif ss == "caps_line":
            cap_style = ParagraphStyle("Caps", fontName="Helvetica", fontSize=s(8.5), textColor=MID_GRAY, spaceBefore=s(6))
            elements.append(Paragraph(title.upper(), cap_style))
            elements.append(Spacer(1, s(1)))
            elements.append(HRFlowable(width="100%", thickness=0.5, color=LIGHT_GRAY))

        elif ss == "bold_bottom_border":
            bb_style = ParagraphStyle("BB", fontName="Helvetica-Bold", fontSize=s(10), textColor=ACCENT, spaceBefore=s(6))
            elements.append(Paragraph(title, bb_style))
            elements.append(Spacer(1, s(1)))
            elements.append(HRFlowable(width="100%", thickness=2.5, color=ACCENT))

        elif ss == "accent_dot":
            dot_color = tpl["accent"]
            dot_html = f'<font color="{dot_color}">&bull;&bull;</font>'
            ad_style = ParagraphStyle("AD", fontName="Helvetica-Bold", fontSize=s(10), textColor=BLACK, spaceBefore=s(6))
            elements.append(Paragraph(f'{dot_html}  {title}', ad_style))
            elements.append(Spacer(1, s(1)))

        elif ss == "soft_bg":
            light_bg = colors.HexColor("#EBF5FB")
            inner = ParagraphStyle("SB", fontName="Helvetica-Bold", fontSize=s(9), textColor=ACCENT, leftIndent=6)
            tbl = Table([[Paragraph(title, inner)]], colWidths=[page_w])
            tbl.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), light_bg),
                ("TOPPADDING", (0, 0), (-1, -1), s(3)),
                ("BOTTOMPADDING", (0, 0), (-1, -1), s(3)),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("LINEBEFORESTARTX", (0, 0), (0, -1), 3, ACCENT),
            ]))
            elements.append(Spacer(1, s(4)))
            elements.append(tbl)

        elif ss == "color_tag":
            tag_style = ParagraphStyle("CT", fontName="Helvetica-Bold", fontSize=s(8.5), textColor=WHITE, leftIndent=4)
            tbl = Table([[Paragraph(title, tag_style)]], colWidths=[page_w])
            tbl.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), ACCENT),
                ("TOPPADDING", (0, 0), (-1, -1), s(2)),
                ("BOTTOMPADDING", (0, 0), (-1, -1), s(2)),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ]))
            elements.append(Spacer(1, s(5)))
            elements.append(tbl)

        elif ss == "plain_bold":
            pb_style = ParagraphStyle("PB", fontName="Helvetica-Bold", fontSize=s(10), textColor=ACCENT, spaceBefore=s(6))
            elements.append(Paragraph(title.upper(), pb_style))

        elif ss == "block_line":
            block_html = f'<font color="{tpl["primary"]}">&#9608;</font>'
            bl_style = ParagraphStyle("BL", fontName="Helvetica-Bold", fontSize=s(10), textColor=BLACK, spaceBefore=s(6))
            elements.append(Paragraph(f'{block_html}  {title}', bl_style))
            elements.append(Spacer(1, s(1)))
            elements.append(HRFlowable(width="100%", thickness=1, color=LIGHT_ACCENT))

        elements.append(Spacer(1, s(1)))
        return elements

    def job_entry(job):
        elements = []
        ds = tpl["date_style"]

        if ds == "underline_below":
            elements.append(Paragraph(job["puesto"], job_title_style))
            cd = job["empresa"]
            if job["fechas"]:
                cd += "  |  " + job["fechas"]
            elements.append(Paragraph(cd, company_style))

        elif ds == "italic_right":
            title_p = Paragraph(job["puesto"], job_title_style)
            date_p = Paragraph(job["fechas"], date_style_r)
            tbl = Table([[title_p, date_p]], colWidths=[page_w * 0.65, page_w * 0.35])
            tbl.setStyle(TableStyle([
                ("VALIGN", (0, 0), (-1, -1), "BOTTOM"),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
            ]))
            elements.append(tbl)
            elements.append(Paragraph(job["empresa"], company_style))

        elif ds == "bold_top":
            date_co_style = ParagraphStyle("DateCo", fontName="Helvetica-Bold", fontSize=s(8.5), textColor=LIGHT_ACCENT, spaceAfter=s(0.5))
            date_text = job["fechas"]
            if job["empresa"]:
                date_text = job["empresa"] + "  |  " + date_text
            elements.append(Paragraph(date_text, date_co_style))
            elements.append(Paragraph(job["puesto"], job_title_style))

        elif ds == "right_gray":
            title_p = Paragraph(f'<b>{job["puesto"]}</b>  -  {job["empresa"]}',
                ParagraphStyle("JT", fontName="Helvetica", fontSize=s(9), textColor=BLACK))
            date_p = Paragraph(job["fechas"], ParagraphStyle("DG", fontName="Helvetica", fontSize=s(8), textColor=LIGHT_GRAY, alignment=TA_RIGHT))
            tbl = Table([[title_p, date_p]], colWidths=[page_w * 0.7, page_w * 0.3])
            tbl.setStyle(TableStyle([
                ("VALIGN", (0, 0), (-1, -1), "BOTTOM"),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), s(1)),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
            ]))
            elements.append(tbl)

        elif ds == "inline_pipe":
            inline_style = ParagraphStyle("IP", fontName="Helvetica-Bold", fontSize=s(9.5), textColor=ACCENT, spaceAfter=s(0.5))
            text = job["puesto"]
            if job["empresa"]:
                text += f'  <font color="{tpl["accent"]}">|</font>  {job["empresa"]}'
            if job["fechas"]:
                text += f'  <font color="{tpl["accent"]}">|</font>  <font size="{s(8)}">{job["fechas"]}</font>'
            elements.append(Paragraph(text, inline_style))

        elif ds == "colored_date":
            elements.append(Paragraph(job["puesto"], job_title_style))
            co_line = job["empresa"]
            if job["fechas"]:
                co_line += f'  |  <font color="{tpl["accent"]}">{job["fechas"]}</font>'
            elements.append(Paragraph(co_line, company_style))

        elif ds == "parens_after":
            title_text = f'<b>{job["puesto"]}</b>'
            if job["empresa"]:
                title_text += f',  {job["empresa"]}'
            if job["fechas"]:
                title_text += f'  <font color="#999999">({job["fechas"]})</font>'
            pa_style = ParagraphStyle("PA", fontName="Helvetica", fontSize=s(9.5), textColor=ACCENT, spaceAfter=s(0.5))
            elements.append(Paragraph(title_text, pa_style))

        elif ds == "dash_sep":
            elements.append(Paragraph(job["puesto"], job_title_style))
            dash_line = job["empresa"]
            if job["fechas"]:
                dash_line += f'  -  {job["fechas"]}'
            elements.append(Paragraph(dash_line, company_style))

        elif ds == "simple_line":
            sl_style = ParagraphStyle("SL", fontName="Helvetica-Bold", fontSize=s(9.5), textColor=ACCENT, spaceAfter=s(0.5))
            elements.append(Paragraph(job["puesto"], sl_style))
            info = job["empresa"]
            if job["fechas"]:
                info += "  |  " + job["fechas"]
            elements.append(Paragraph(info, ParagraphStyle("SI", fontName="Helvetica", fontSize=s(8.5), textColor=MID_GRAY, spaceAfter=s(0.5))))

        elif ds == "bold_date_left":
            elements.append(Paragraph(job["puesto"], job_title_style))
            co_date = job["empresa"]
            if job["fechas"]:
                co_date += f'  |  <font color="{tpl["accent"]}"><b>{job["fechas"]}</b></font>'
            elements.append(Paragraph(co_date, company_style))

        for b in job["bullets"]:
            elements.append(Paragraph(f"&bull; {b}", bullet_style))
        elements.append(Spacer(1, s(1.5)))
        return elements

    story = []
    if tpl.get("dark_header"):
        name_dark = ParagraphStyle("NameDK", fontName="Helvetica-Bold", fontSize=s(20), leading=s(24), textColor=WHITE, spaceAfter=0)
        contact_dark = ParagraphStyle("ContactDK", fontName="Helvetica", fontSize=s(8.5), textColor=colors.HexColor("#CCDDEE"), spaceAfter=0)
        subtitle_dark = ParagraphStyle("SubDK", fontName="Helvetica-Oblique", fontSize=s(10), textColor=colors.HexColor("#88BBDD"), spaceAfter=0)
        hdr_story = [Paragraph(data["nombre"], name_dark), Spacer(1, s(2)), Paragraph(data["contacto"], contact_dark)]
        if data["titulo"]:
            hdr_story += [Spacer(1, s(2)), Paragraph(data["titulo"], subtitle_dark)]
        hdr_tbl = Table([[hdr_story]], colWidths=[page_w])
        hdr_tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), ACCENT),
            ("TOPPADDING", (0, 0), (-1, -1), s(8)),
            ("BOTTOMPADDING", (0, 0), (-1, -1), s(8)),
            ("LEFTPADDING", (0, 0), (-1, -1), 10),
        ]))
        story.append(hdr_tbl)
    else:
        story.append(Paragraph(data["nombre"], name_style))
        story.append(Spacer(1, s(2)))
        story.append(Paragraph(data["contacto"], contact_style))
        if data["titulo"]:
            story.append(Spacer(1, s(2)))
            story.append(Paragraph(data["titulo"], subtitle_style))
    story.append(Spacer(1, s(4)))

    sep = tpl["separator"]
    if sep == "hr":
        story.append(HRFlowable(width="100%", thickness=1.5, color=ACCENT))
    elif sep == "double_line":
        story.append(HRFlowable(width="100%", thickness=0.5, color=ACCENT))
        story.append(Spacer(1, 2))
        story.append(HRFlowable(width="100%", thickness=0.5, color=ACCENT))
    elif sep == "accent_thick":
        story.append(HRFlowable(width="100%", thickness=3, color=ACCENT))
    elif sep == "thin_line":
        story.append(HRFlowable(width="100%", thickness=0.5, color=LIGHT_GRAY))
    story.append(Spacer(1, s(2)))

    if data["resumen"]:
        for el in section_header(labels["resumen"]):
            story.append(el)
        story.append(Paragraph(data["resumen"], body_style))
        story.append(Spacer(1, s(2)))

    if data["logros"]:
        for el in section_header(labels["logros"]):
            story.append(el)
        for logro in data["logros"]:
            story.append(Paragraph(f"&bull; {logro}", bullet_style))
        story.append(Spacer(1, s(2)))

    if data["experiencia"]:
        for el in section_header(labels["experiencia"]):
            story.append(el)
        for job in data["experiencia"]:
            for el in job_entry(job):
                story.append(el)

    if data["educacion"]:
        for el in section_header(labels["educacion"]):
            story.append(el)
        for edu in data["educacion"]:
            story.append(Paragraph(edu["titulo"], job_title_style))
            story.append(Paragraph(edu["institucion"], company_style))
            story.append(Spacer(1, s(1)))

    if data["cursos"]:
        for el in section_header(labels["cursos"]):
            story.append(el)
        for c in data["cursos"]:
            story.append(Paragraph(f"&bull; {c}", bullet_style))
        story.append(Spacer(1, s(2)))

    if data["habilidades"]:
        for el in section_header(labels["habilidades"]):
            story.append(el)
        rows = []
        for label, value in data["habilidades"]:
            rows.append([Paragraph(label, skill_label_s), Paragraph(value, skill_value_s)])
        tbl = Table(rows, colWidths=[2.0 * inch, page_w - 2.0 * inch])
        tbl.setStyle(TableStyle([
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("TOPPADDING", (0, 0), (-1, -1), 0),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
        ]))
        story.append(tbl)
        story.append(Spacer(1, s(2)))

    if data["idiomas"]:
        for el in section_header(labels["idiomas"]):
            story.append(el)
        for idioma in data["idiomas"]:
            story.append(Paragraph(f"&bull; {idioma}", bullet_style))

    return story


def generate_pdf(data, output_path, labels=None):
    if labels is None:
        labels = LABELS_ES
    from reportlab.lib.pagesizes import letter

    class PageCounter(SimpleDocTemplate):
        def __init__(self, *args, **kwargs):
            super().__init__(*args, **kwargs)
            self.page_count = 0
        def afterPage(self):
            self.page_count += 1

    best_scale = 1.0
    for attempt in range(20):
        sc = 1.0 - (attempt * 0.03)
        if sc < 0.55:
            break
        buf = io.BytesIO()
        counter = PageCounter(buf, pagesize=letter, rightMargin=0.5*inch, leftMargin=0.5*inch, topMargin=0.35*inch, bottomMargin=0.35*inch)
        story = _build_pdf_story(data, sc, labels)
        counter.build(story)
        if counter.page_count <= 1:
            best_scale = sc
            break
        best_scale = sc

    story = _build_pdf_story(data, best_scale, labels)
    doc = SimpleDocTemplate(output_path, pagesize=letter, rightMargin=0.5*inch, leftMargin=0.5*inch, topMargin=0.35*inch, bottomMargin=0.35*inch)
    doc.build(story)

    pct = int(best_scale * 100)
    if pct < 100:
        print(f"  PDF generado:  {output_path}  (escala: {pct}%)")
    else:
        print(f"  PDF generado:  {output_path}")
    return best_scale


def generate_word(data, output_path, scale=1.0, labels=None):
    if labels is None:
        labels = LABELS_ES

    tpl = TEMPLATES[data["theme"]]

    def hex_rgb(h):
        h = h.lstrip("#")
        return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

    C_PRIMARY = hex_rgb(tpl["primary"])
    C_ACCENT = hex_rgb(tpl["accent"])
    C_TEXT = RGBColor(0x33, 0x33, 0x33)
    C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    C_GRAY = RGBColor(0x99, 0x99, 0x99)
    C_BLACK = RGBColor(0x22, 0x22, 0x22)
    BG_HEX = tpl["primary"].lstrip("#")
    ACCENT_HEX = tpl["accent"].lstrip("#")

    def s(base):
        return base * scale

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(0.8)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(s(9))
    style.font.color.rgb = C_TEXT
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = Pt(s(11))

    header_align = WD_ALIGN_PARAGRAPH.CENTER if tpl["header_align"] == "center" else WD_ALIGN_PARAGRAPH.LEFT

    def sp(p, before=0, after=0, line=11):
        pf = p.paragraph_format
        pf.space_before = Pt(s(before))
        pf.space_after = Pt(s(after))
        pf.line_spacing = Pt(s(line))

    def no_borders(table):
        tbl = table._tbl
        tblPr = tbl.tblPr
        tblPr.append(parse_xml(
            f'<w:tblBorders {nsdecls("w")}> '
            f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'</w:tblBorders>'
        ))
        tblPr.append(parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>'))

    def cell_margins(cell, top=0, left=40, bottom=0, right=40):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcPr.append(parse_xml(
            f'<w:tcMar {nsdecls("w")}> '
            f'  <w:top w:w="{top}" w:type="dxa"/>'
            f'  <w:left w:w="{left}" w:type="dxa"/>'
            f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
            f'  <w:right w:w="{right}" w:type="dxa"/>'
            f'</w:tcMar>'
        ))

    def add_bullet(text):
        p = doc.add_paragraph()
        sp(p, 0, 0, 10.5)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.first_line_indent = Cm(-0.3)
        run = p.add_run("- " + text)
        run.font.size = Pt(s(9))
        run.font.color.rgb = C_TEXT

    def add_section_header(title):
        ss = tpl["section_style"]

        if ss == "bar":
            table = doc.add_table(rows=1, cols=1)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell = table.cell(0, 0)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{BG_HEX}"/>'))
            cell_margins(cell, 15, 80, 15, 80)
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(title)
            run.bold = True
            run.font.color.rgb = C_WHITE
            run.font.size = Pt(s(9.5))
            run.font.name = 'Calibri'
            sp(p, 0, 0, 12)
            no_borders(table)

        elif ss == "text_underline":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(title)
            run.bold = True
            run.font.size = Pt(s(10))
            run.font.color.rgb = C_PRIMARY
            sp(p, 6, 0, 12)
            p2 = doc.add_paragraph()
            p2.paragraph_format.space_before = Pt(0)
            p2.paragraph_format.space_after = Pt(s(2))
            pPr = p2._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}> '
                f'  <w:bottom w:val="single" w:sz="8" w:space="1" w:color="{BG_HEX}"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)

        elif ss == "left_border":
            table = doc.add_table(rows=1, cols=1)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell = table.cell(0, 0)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F8FC"/>'))
            tcBorders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}> '
                f'  <w:left w:val="single" w:sz="18" w:space="0" w:color="{BG_HEX}"/>'
                f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'</w:tcBorders>'
            )
            tcPr.append(tcBorders)
            cell_margins(cell, 30, 120, 30, 80)
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(title)
            run.bold = True
            run.font.size = Pt(s(9.5))
            run.font.color.rgb = C_PRIMARY
            sp(p, 0, 0, 12)
            tblPr = table._tbl.tblPr
            tblPr.append(parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>'))
            borders = parse_xml(
                f'<w:tblBorders {nsdecls("w")}> '
                f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'</w:tblBorders>'
            )
            tblPr.append(borders)

        elif ss == "caps_line":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(title.upper())
            run.font.size = Pt(s(8.5))
            run.font.color.rgb = C_GRAY
            run.font.name = 'Calibri'
            sp(p, 6, 0, 11)
            p2 = doc.add_paragraph()
            sp(p2, 0, s(1), 2)
            pPr = p2._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}> '
                f'  <w:bottom w:val="single" w:sz="2" w:space="1" w:color="999999"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)

        elif ss == "bold_bottom_border":
            p = doc.add_paragraph()
            run = p.add_run(title)
            run.bold = True
            run.font.size = Pt(s(10))
            run.font.color.rgb = C_PRIMARY
            sp(p, 6, 0, 12)
            p2 = doc.add_paragraph()
            sp(p2, 0, s(1), 2)
            pPr = p2._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}> '
                f'  <w:bottom w:val="single" w:sz="16" w:space="1" w:color="{BG_HEX}"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)

        elif ss == "accent_dot":
            p = doc.add_paragraph()
            sp(p, 6, 1, 12)
            run_dot = p.add_run("**  ")
            run_dot.font.size = Pt(s(8))
            run_dot.font.color.rgb = C_ACCENT
            run_text = p.add_run(title)
            run_text.bold = True
            run_text.font.size = Pt(s(10))
            run_text.font.color.rgb = C_BLACK

        elif ss == "soft_bg":
            light_hex = "EBF5FB"
            table = doc.add_table(rows=1, cols=1)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell = table.cell(0, 0)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{light_hex}"/>'))
            tcBorders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}> '
                f'  <w:left w:val="single" w:sz="18" w:space="0" w:color="{BG_HEX}"/>'
                f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'</w:tcBorders>'
            )
            tcPr.append(tcBorders)
            cell_margins(cell, 25, 100, 25, 80)
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(title)
            run.bold = True
            run.font.size = Pt(s(9))
            run.font.color.rgb = C_PRIMARY
            sp(p, 0, 0, 12)
            tblPr = table._tbl.tblPr
            tblPr.append(parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>'))
            borders = parse_xml(
                f'<w:tblBorders {nsdecls("w")}> '
                f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'</w:tblBorders>'
            )
            tblPr.append(borders)

        elif ss == "color_tag":
            table = doc.add_table(rows=1, cols=1)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell = table.cell(0, 0)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{ACCENT_HEX}"/>'))
            cell_margins(cell, 15, 80, 15, 80)
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(title)
            run.bold = True
            run.font.color.rgb = C_WHITE
            run.font.size = Pt(s(8.5))
            sp(p, 0, 0, 11)
            no_borders(table)

        elif ss == "plain_bold":
            p = doc.add_paragraph()
            run = p.add_run(title.upper())
            run.bold = True
            run.font.size = Pt(s(10))
            run.font.color.rgb = C_PRIMARY
            sp(p, 6, 1, 12)

        elif ss == "block_line":
            p = doc.add_paragraph()
            sp(p, 6, 0, 12)
            run_block = p.add_run("█  ")
            run_block.font.size = Pt(s(9))
            run_block.font.color.rgb = C_PRIMARY
            run_text = p.add_run(title)
            run_text.bold = True
            run_text.font.size = Pt(s(10))
            run_text.font.color.rgb = C_BLACK
            p2 = doc.add_paragraph()
            sp(p2, 0, s(1), 2)
            pPr = p2._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}> '
                f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="{ACCENT_HEX}"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)

    def add_job_entry(job):
        ds = tpl["date_style"]

        if ds == "underline_below":
            p = doc.add_paragraph()
            sp(p, 2, 0, 11)
            cd = job["empresa"]
            if job["fechas"]:
                cd += " | " + job["fechas"]
            run = p.add_run(cd)
            run.font.color.rgb = C_ACCENT
            run.font.size = Pt(s(9))
            run.underline = True
            p = doc.add_paragraph()
            sp(p, 0, 0, 11)
            run = p.add_run(job["puesto"])
            run.bold = True
            run.font.size = Pt(s(9))
            run.font.color.rgb = C_TEXT

        elif ds == "italic_right":
            table = doc.add_table(rows=1, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            no_borders(table)
            c1 = table.cell(0, 0)
            c1.paragraphs[0].clear()
            run = c1.paragraphs[0].add_run(job["puesto"])
            run.bold = True
            run.font.size = Pt(s(9.5))
            run.font.color.rgb = C_PRIMARY
            sp(c1.paragraphs[0], 2, 0, 11)
            c2 = table.cell(0, 1)
            c2.paragraphs[0].clear()
            c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run2 = c2.paragraphs[0].add_run(job["fechas"])
            run2.italic = True
            run2.font.size = Pt(s(8.5))
            run2.font.color.rgb = C_GRAY
            sp(c2.paragraphs[0], 2, 0, 11)
            p = doc.add_paragraph()
            sp(p, 0, 0, 11)
            run = p.add_run(job["empresa"])
            run.font.size = Pt(s(9))
            run.font.color.rgb = C_ACCENT

        elif ds == "bold_top":
            p = doc.add_paragraph()
            sp(p, 2, 0, 11)
            dt = job["empresa"]
            if job["fechas"]:
                dt += "  |  " + job["fechas"]
            run = p.add_run(dt)
            run.bold = True
            run.font.size = Pt(s(8.5))
            run.font.color.rgb = C_ACCENT
            p = doc.add_paragraph()
            sp(p, 0, 0, 11)
            run = p.add_run(job["puesto"])
            run.bold = True
            run.font.size = Pt(s(9.5))
            run.font.color.rgb = C_PRIMARY

        elif ds == "right_gray":
            table = doc.add_table(rows=1, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            no_borders(table)
            c1 = table.cell(0, 0)
            c1.paragraphs[0].clear()
            run1 = c1.paragraphs[0].add_run(job["puesto"])
            run1.bold = True
            run1.font.size = Pt(s(9))
            run1.font.color.rgb = C_BLACK
            run1b = c1.paragraphs[0].add_run("  -  " + job["empresa"])
            run1b.font.size = Pt(s(9))
            run1b.font.color.rgb = C_TEXT
            sp(c1.paragraphs[0], 2, 0, 11)
            c2 = table.cell(0, 1)
            c2.paragraphs[0].clear()
            c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run2 = c2.paragraphs[0].add_run(job["fechas"])
            run2.font.size = Pt(s(8))
            run2.font.color.rgb = C_GRAY
            sp(c2.paragraphs[0], 2, 0, 11)

        elif ds == "inline_pipe":
            p = doc.add_paragraph()
            sp(p, 2, 0, 11)
            run = p.add_run(job["puesto"])
            run.bold = True
            run.font.size = Pt(s(9.5))
            run.font.color.rgb = C_PRIMARY
            if job["empresa"]:
                run2 = p.add_run("  |  " + job["empresa"])
                run2.font.size = Pt(s(9))
                run2.font.color.rgb = C_ACCENT
            if job["fechas"]:
                run3 = p.add_run("  |  " + job["fechas"])
                run3.font.size = Pt(s(8))
                run3.font.color.rgb = C_TEXT

        elif ds == "colored_date":
            p = doc.add_paragraph()
            sp(p, 2, 0, 11)
            run = p.add_run(job["puesto"])
            run.bold = True
            run.font.size = Pt(s(9.5))
            run.font.color.rgb = C_PRIMARY
            p = doc.add_paragraph()
            sp(p, 0, 0, 11)
            run = p.add_run(job["empresa"])
            run.font.size = Pt(s(9))
            run.font.color.rgb = C_TEXT
            if job["fechas"]:
                run2 = p.add_run("  |  " + job["fechas"])
                run2.font.size = Pt(s(9))
                run2.font.color.rgb = C_ACCENT

        elif ds == "parens_after":
            p = doc.add_paragraph()
            sp(p, 2, 0, 11)
            run = p.add_run(job["puesto"])
            run.bold = True
            run.font.size = Pt(s(9.5))
            run.font.color.rgb = C_PRIMARY
            if job["empresa"]:
                run2 = p.add_run(",  " + job["empresa"])
                run2.font.size = Pt(s(9))
                run2.font.color.rgb = C_TEXT
            if job["fechas"]:
                run3 = p.add_run("  (" + job["fechas"] + ")")
                run3.font.size = Pt(s(8.5))
                run3.font.color.rgb = C_GRAY

        elif ds == "dash_sep":
            p = doc.add_paragraph()
            sp(p, 2, 0, 11)
            run = p.add_run(job["puesto"])
            run.bold = True
            run.font.size = Pt(s(9.5))
            run.font.color.rgb = C_PRIMARY
            p = doc.add_paragraph()
            sp(p, 0, 0, 11)
            dl = job["empresa"]
            if job["fechas"]:
                dl += "  -  " + job["fechas"]
            run = p.add_run(dl)
            run.font.size = Pt(s(9))
            run.font.color.rgb = C_ACCENT

        elif ds == "simple_line":
            p = doc.add_paragraph()
            sp(p, 2, 0, 11)
            run = p.add_run(job["puesto"])
            run.bold = True
            run.font.size = Pt(s(9.5))
            run.font.color.rgb = C_PRIMARY
            p = doc.add_paragraph()
            sp(p, 0, 0, 11)
            info = job["empresa"]
            if job["fechas"]:
                info += "  |  " + job["fechas"]
            run = p.add_run(info)
            run.font.size = Pt(s(8.5))
            run.font.color.rgb = C_GRAY

        elif ds == "bold_date_left":
            p = doc.add_paragraph()
            sp(p, 2, 0, 11)
            run = p.add_run(job["puesto"])
            run.bold = True
            run.font.size = Pt(s(9.5))
            run.font.color.rgb = C_PRIMARY
            p = doc.add_paragraph()
            sp(p, 0, 0, 11)
            run = p.add_run(job["empresa"])
            run.font.size = Pt(s(9))
            run.font.color.rgb = C_ACCENT
            if job["fechas"]:
                run2 = p.add_run("  |  " + job["fechas"])
                run2.bold = True
                run2.font.size = Pt(s(8.5))
                run2.font.color.rgb = C_ACCENT

        for b in job["bullets"]:
            add_bullet(b)

    if tpl.get("dark_header"):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{BG_HEX}"/>'))
        cell_margins(cell, 60, 100, 60, 100)
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(data["nombre"])
        run.bold = True
        run.font.size = Pt(s(22))
        run.font.color.rgb = C_WHITE
        sp(p, 0, 2, 24)
        p2 = cell.add_paragraph()
        run2 = p2.add_run(data["contacto"])
        run2.font.size = Pt(s(9))
        run2.font.color.rgb = RGBColor(0xCC, 0xDD, 0xEE)
        sp(p2, 0, 0, 11)
        if data["titulo"]:
            p3 = cell.add_paragraph()
            run3 = p3.add_run(data["titulo"])
            run3.font.size = Pt(s(10))
            run3.font.color.rgb = RGBColor(0x88, 0xBB, 0xDD)
            run3.italic = True
            sp(p3, 1, 0, 12)
        tblPr = table._tbl.tblPr
        tblPr.append(parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>'))
        brd = parse_xml(
            f'<w:tblBorders {nsdecls("w")}> '
            f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'</w:tblBorders>'
        )
        tblPr.append(brd)
    else:
        p = doc.add_paragraph()
        p.alignment = header_align
        run = p.add_run(data["nombre"])
        run.bold = True
        run.font.size = Pt(s(22))
        run.font.color.rgb = C_PRIMARY
        sp(p, 0, 0, 24)

        p = doc.add_paragraph()
        p.alignment = header_align
        run = p.add_run(data["contacto"])
        run.font.size = Pt(s(9))
        run.font.color.rgb = C_TEXT
        sp(p, 0, 0, 11)

        if data["titulo"]:
            p = doc.add_paragraph()
            p.alignment = header_align
            run = p.add_run(data["titulo"])
            run.font.size = Pt(s(10))
            run.font.color.rgb = C_ACCENT
            run.italic = True
            sp(p, 0, 2, 12)

    sep = tpl["separator"]
    if sep in ("hr", "accent_thick", "thin_line"):
        p = doc.add_paragraph()
        sp(p, 0, s(2), 2)
        pPr = p._p.get_or_add_pPr()
        thickness = "12" if sep == "hr" else ("20" if sep == "accent_thick" else "2")
        color_hex = BG_HEX if sep != "thin_line" else "999999"
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}> '
            f'  <w:bottom w:val="single" w:sz="{thickness}" w:space="1" w:color="{color_hex}"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)
    elif sep == "double_line":
        for _ in range(2):
            p = doc.add_paragraph()
            sp(p, 0, 1, 2)
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}> '
                f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="{BG_HEX}"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)

    if data["resumen"]:
        add_section_header(labels["resumen"])
        p = doc.add_paragraph()
        run = p.add_run(data["resumen"])
        run.font.size = Pt(s(9))
        run.font.color.rgb = C_TEXT
        sp(p, 1, 2, 10.5)

    if data["logros"]:
        add_section_header(labels["logros"])
        for logro in data["logros"]:
            add_bullet(logro)

    if data["experiencia"]:
        add_section_header(labels["experiencia"])
        for job in data["experiencia"]:
            add_job_entry(job)

    if data["educacion"]:
        add_section_header(labels["educacion"])
        for edu in data["educacion"]:
            p = doc.add_paragraph()
            run = p.add_run(edu["titulo"])
            run.bold = True
            run.font.size = Pt(s(9))
            run.font.color.rgb = C_TEXT
            sp(p, 1, 0, 11)
            p = doc.add_paragraph()
            run = p.add_run(edu["institucion"])
            run.font.size = Pt(s(9))
            run.font.color.rgb = C_ACCENT
            sp(p, 0, 1, 11)

    if data["cursos"]:
        add_section_header(labels["cursos"])
        for c in data["cursos"]:
            add_bullet(c)

    if data["habilidades"]:
        add_section_header(labels["habilidades"])
        for cat, detalle in data["habilidades"]:
            p = doc.add_paragraph()
            sp(p, 0, 0, 10.5)
            p.paragraph_format.left_indent = Cm(0.3)
            run_b = p.add_run(cat)
            run_b.bold = True
            run_b.font.size = Pt(s(9))
            run_b.font.color.rgb = C_TEXT
            run_r = p.add_run(" " + detalle)
            run_r.font.size = Pt(s(9))
            run_r.font.color.rgb = C_TEXT

    if data["idiomas"]:
        add_section_header(labels["idiomas"])
        for idioma in data["idiomas"]:
            add_bullet(idioma)

    doc.save(output_path)
    pct = int(scale * 100)
    if pct < 100:
        print(f"  Word generado: {output_path}  (escala: {pct}%)")
    else:
        print(f"  Word generado: {output_path}")
